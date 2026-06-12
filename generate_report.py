"""
棋评 Word 文档生成器 v4
严格按照微信公众号"深蓝棋评"参考文章风格：
  - 内联 PGN（如 "1. e4 e5 2. Nf3 Nc6 3. Bb5 a6 4. Bxc6 dxc6，棋局进入..."）
  - 每 3~5 步一组，PGN 在前、讲解在后
  - 每个 PGN 段配棋盘图示
  - "..." 作为续弈标记
  - 关键局面展开详析（"若...将..." / "此处引擎认为..."）
  - 完整 PGN 附于文末

用法: python generate_report.py [--output 输出路径.docx] [--no-images]
"""

import sys, json, re, os, tempfile
from pathlib import Path
from datetime import datetime

sys.stdout.reconfigure(encoding="utf-8")

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("请安装 python-docx: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    print("请安装 Pillow: pip install Pillow")
    sys.exit(1)

import chess, chess.pgn


# ===================== 样式常量 =====================
FONT = '微软雅黑'
C_TITLE   = RGBColor(25, 55, 100)
C_BODY    = RGBColor(35, 35, 40)
C_FOOTER  = RGBColor(140, 140, 150)
C_ENGINE  = RGBColor(80, 80, 90)
C_RED     = RGBColor(190, 40, 40)


def add_p(doc, text, size=11, bold=False, color=None, align=None,
          after=8, indent=True):
    """统一添加段落"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent and align is None:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    run.bold = bold
    run.font.color.rgb = (color or C_BODY)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.65
    return p


def sep(doc):
    add_p(doc, "· · ·", size=10, color=C_FOOTER, align=WD_ALIGN_PARAGRAPH.CENTER,
          after=10, indent=False)


# ===================== 工具 =====================

def round_num(move_number: int) -> int:
    """走棋编号 → 回合编号"""
    return (move_number + 1) // 2


def format_pgn_segment(sans: list, start_move: int) -> str:
    """
    将一段 SAN 列表格式化为内联 PGN。
    如 ['e4','e5','Nf3','Nc6'] start_move=1 → "1. e4 e5 2. Nf3 Nc6"
    """
    parts = []
    for i, san in enumerate(sans):
        mv = start_move + i
        if mv % 2 == 1:  # 白方走棋，加编号
            r = round_num(mv)
            parts.append(f"{r}. {san}")
        else:
            parts.append(san)
    return " ".join(parts)


def clean_commentary(text: str) -> str:
    """移除画面指令标签和 Markdown 格式"""
    for tag in ['高亮', '威胁', '选中', '箭头']:
        text = re.sub(rf'\[{tag}\s*[^\]]+\]', '', text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    return text.strip()


# ===================== 棋盘图片渲染 =====================

class BoardImageRenderer:
    """紧凑棋盘渲染器 — 用于 Word 文档嵌入"""

    BOARD_SIZE = 360   # 总像素
    SQ = BOARD_SIZE // 8

    LIGHT = (240, 217, 181)
    DARK  = (181, 136, 99)
    LAST_MOVE_HL = (255, 255, 0, 90)

    def __init__(self, pieces_dir: Path):
        self.piece_imgs = {}
        mapping = {
            'K': 'wK', 'Q': 'wQ', 'R': 'wR', 'B': 'wB', 'N': 'wN', 'P': 'wP',
            'k': 'bK', 'q': 'bQ', 'r': 'bR', 'b': 'bB', 'n': 'bN', 'p': 'bP',
        }
        for symbol, fname in mapping.items():
            fp = pieces_dir / f"{fname}.png"
            if fp.exists():
                img = Image.open(fp).convert('RGBA')
                # 棋子占格子 92%，比之前更大更清晰
                target = int(self.SQ * 0.92)
                if img.width != target:
                    img = img.resize((target, target), Image.LANCZOS)
                self.piece_imgs[symbol] = img

    def render(self, board: chess.Board, last_from=None, last_to=None,
               highlight_squares=None, annotation=None) -> Image.Image:
        """渲染棋盘为 PIL Image（可选着法标注）"""
        img = Image.new('RGBA', (self.BOARD_SIZE, self.BOARD_SIZE), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)

        # 格子 — a1(0,0) 必须是深色格
        for rank in range(8):
            for file in range(8):
                x = file * self.SQ
                y = (7 - rank) * self.SQ
                # a1=(file=0,rank=0): (0+0)%2=0 → DARK ✓
                color = self.DARK if (file + rank) % 2 == 0 else self.LIGHT
                draw.rectangle([x, y, x + self.SQ - 1, y + self.SQ - 1], fill=color)

        # 高亮上一步走棋格
        if last_from is not None:
            for sq in [last_from, last_to]:
                if sq is None:
                    continue
                f = chess.square_file(sq)
                r = chess.square_rank(sq)
                x = f * self.SQ
                y = (7 - r) * self.SQ
                overlay = Image.new('RGBA', (self.SQ, self.SQ), self.LAST_MOVE_HL)
                img.paste(overlay, (x, y), overlay)

        # 额外高亮
        if highlight_squares:
            for sq in highlight_squares:
                f = chess.square_file(sq)
                r = chess.square_rank(sq)
                x = f * self.SQ
                y = (7 - r) * self.SQ
                overlay = Image.new('RGBA', (self.SQ, self.SQ), (255, 80, 50, 70))
                img.paste(overlay, (x, y), overlay)

        # 棋子
        for square in chess.SQUARES:
            piece = board.piece_at(square)
            if not piece:
                continue
            symbol = piece.symbol()
            p_img = self.piece_imgs.get(symbol)
            if p_img is None:
                continue
            f = chess.square_file(square)
            r = chess.square_rank(square)
            px = f * self.SQ + (self.SQ - p_img.width) // 2
            py = (7 - r) * self.SQ + (self.SQ - p_img.height) // 2
            img.paste(p_img, (px, py), p_img)

        # 坐标标注
        try:
            font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 8)
        except Exception:
            font = ImageFont.load_default()

        for file in range(8):
            x = file * self.SQ + self.SQ // 2 - 3
            y = self.BOARD_SIZE - 9
            draw.text((x, y), chr(ord('a') + file), fill=(100, 100, 100), font=font)
        for rank in range(8):
            x = 2
            y = (7 - rank) * self.SQ + self.SQ // 2 - 5
            draw.text((x, y), str(rank + 1), fill=(100, 100, 100), font=font)

        # 着法标注
        if annotation and last_to is not None:
            file = chess.square_file(last_to)
            rank = chess.square_rank(last_to)
            sq_x = file * self.SQ
            sq_y = (7 - rank) * self.SQ
            badge = int(self.SQ * 0.32)
            margin = int(self.SQ * 0.04)
            bx = sq_x + self.SQ - badge + margin
            by = sq_y - margin
            draw.ellipse([bx, by, bx + badge, by + badge],
                         fill=(200, 50, 50), outline=(160, 30, 30), width=2)
            try:
                af = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", int(badge * 0.6))
            except Exception:
                af = ImageFont.load_default()
            text = annotation if annotation in ("?", "?!", "??") else "?"
            tb = draw.textbbox((0, 0), text, font=af)
            tw, th = tb[2] - tb[0], tb[3] - tb[1]
            draw.text((bx + (badge - tw) // 2, by + (badge - th) // 2 - 1),
                      text, fill=(255, 255, 255), font=af)

        # 边框
        draw.rectangle([0, 0, self.BOARD_SIZE - 1, self.BOARD_SIZE - 1],
                       outline=(50, 50, 50), width=2)

        return img


def replay_to_move(sans: list, move_num: int) -> tuple:
    """重放到第 move_num 步，返回 (board, from_sq, to_sq)"""
    board = chess.Board()
    last_from = None
    last_to = None
    for i, san in enumerate(sans):
        if i >= move_num:
            break
        move = board.parse_san(san)
        last_from = move.from_square
        last_to = move.to_square
        board.push(move)
    return board, last_from, last_to


def find_file(script_dir: Path, filename: str) -> Path:
    """查找文件：先在根目录找，再在 output 子目录找"""
    f = script_dir / filename
    if f.exists():
        return f
    # 搜索 output 子目录
    outdir = script_dir / "output"
    if outdir.exists():
        for sub in sorted(outdir.iterdir(), reverse=True):
            if sub.is_dir():
                f = sub / filename
                if f.exists():
                    return f
    return script_dir / filename  # 返回默认路径（可能不存在）


def load_all(script_dir: Path):
    """加载所有数据"""
    # PGN — 从根目录或 output 找
    pgns = list(script_dir.glob("lichess_pgn*.pgn"))
    pgn_path = pgns[0] if pgns else None

    boards, sans, headers = [], [], {}
    if pgn_path and pgn_path.exists():
        with pgn_path.open("r", encoding="utf-8") as f:
            game = chess.pgn.read_game(f)
        headers = game.headers
        board = game.board()
        for move in game.mainline_moves():
            sans.append(board.san(move))
            board.push(move)
            boards.append(board.copy())

    # 如果根目录没有 PGN，尝试从 output 目录的 merged 文件中提取
    if not sans:
        merged_file = find_file(script_dir, "merged_analysis_commentary.json")
        if merged_file.exists():
            with merged_file.open("r", encoding="utf-8") as f:
                mdata = json.load(f)
            # 尝试提取 PGN 信息
            if "pgn_text" in mdata:
                # 从 PGN 文本重建
                try:
                    import io
                    game = chess.pgn.read_game(io.StringIO(mdata["pgn_text"]))
                    if game:
                        headers = game.headers
                        board = game.board()
                        for move in game.mainline_moves():
                            sans.append(board.san(move))
                            board.push(move)
                            boards.append(board.copy())
                except Exception:
                    pass

    # Analysis JSON
    analysis = find_file(script_dir, "analysis_result.json")
    steps, phases, tb_results = [], [], []
    if analysis.exists():
        with analysis.open("r", encoding="utf-8") as f:
            data = json.load(f)
        steps = data.get("steps", data)
        phases = data.get("phases", [])
        tb_results = data.get("tablebase_results", [])

    # Commentary
    commentary = ""
    step_comms = {}
    comm_file = find_file(script_dir, "commentary.txt")
    if comm_file.exists():
        commentary = comm_file.read_text(encoding="utf-8")
        for m in re.finditer(r"\[STEP (\d+)\]\s*(.*?)(?=\[STEP \d+\]|$)",
                              commentary, re.DOTALL):
            step_comms[int(m.group(1))] = clean_commentary(m.group(2))

    return sans, headers, steps, phases, tb_results, step_comms


# ===================== 文章构建 =====================

def build_article(sans, headers, steps, phases, tb_results, step_comms):
    """
    构建微信公众号风格的讲解文章。
    返回 (title, paragraphs)
    """
    white = headers.get("White", "白方")
    black = headers.get("Black", "黑方")
    opening = headers.get("Opening", "未知开局")
    eco = headers.get("ECO", "?")
    result_str = headers.get("Result", "*")
    game_date = headers.get("UTCDate", "").replace(".", "/")

    result_cn = {"1-0": f"{white} 获胜", "0-1": f"{black} 获胜",
                 "1/2-1/2": "双方和棋"}.get(result_str, result_str)

    title = f"{white} vs {black} — {opening} ({eco})"
    paragraphs = []

    # ===== 简介 =====
    intro = (
        f"今天为大家带来的是{f' {game_date} ' if game_date else ''}的一盘棋，"
        f"由{white}执白对弈{black}。"
        f"棋局进入了{opening}。"
        f"最终{result_cn}。"
        f"下面让我们来对这盘棋进行复盘。"
    )
    paragraphs.append(("intro", intro))

    if not sans:
        paragraphs.append(("body", "（无走棋数据）"))
        return title, paragraphs

    # ===== 按步数分段（每 3~6 步一组） =====
    SEGMENT_SIZE = 4
    total_moves = len(sans)
    i = 0

    # 收集已在正文中覆盖的步号（避免重复输出）
    covered_in_body = set()

    while i < total_moves:
        seg_end = min(i + SEGMENT_SIZE, total_moves)
        seg_sans = sans[i:seg_end]
        seg_start_mv = i + 1

        # PGN 行
        seg_label = "..." if i > 0 else ""
        seg_text = (seg_label + " " + format_pgn_segment(seg_sans, seg_start_mv)
                    if seg_label else format_pgn_segment(seg_sans, seg_start_mv))

        # 讲解文本 — 收集这段的 commentary
        explanations = []
        for mv in range(seg_start_mv, seg_end + 1):
            if mv in step_comms and step_comms[mv]:
                comm = step_comms[mv]
                if len(comm) > 10:
                    explanations.append(comm)
                    covered_in_body.add(mv)

        if explanations:
            seg_text += "，" + "。".join(explanations[:4]) + "。"
            # 清理多余的句号
            seg_text = re.sub(r'。+', '。', seg_text)
            seg_text = re.sub(r'，。', '，', seg_text)

        paragraphs.append(("body", seg_text, seg_end))

        # 只在正文 commentary 未覆盖时，才单独输出深度分析段落
        for mv in range(seg_start_mv, seg_end + 1):
            if mv > len(steps) or mv in covered_in_body:
                continue  # 已在正文 commentary 中提及，不重复

            step = steps[mv - 1]
            s_quality = step.get("quality", "正常")
            s_diff = step.get("score_diff", 0)
            s_cv = step.get("cross_validation")

            # 仅对 commentary 未覆盖的大错输出单独段落
            if s_quality in ("错误", "大错"):
                crit_text = (
                    f"第{round_num(mv)}回合{step['side']}的{step['move_san']}"
                    f"是一步{s_quality}（评分变化 {s_diff:+.1f}）。"
                )
                paragraphs.append(("critical", crit_text))

            # 引擎分歧（仅在未在正文中提及时输出）
            if s_cv and s_cv.get("disagreement_type", "agree") != "agree":
                dtype = s_cv["disagreement_type"]
                sf_score = s_cv.get("stockfish_score", 0)
                lc0_score = s_cv.get("lc0_score", 0)
                if dtype == "disagree_strong":
                    engine_text = (
                        f"有意思的是，引擎在此处看法并不一致："
                        f"Stockfish评分为{sf_score:+.1f}，"
                        f"Lc0神经网络评分为{lc0_score:+.1f}。"
                        f"这往往是静态评价与动态补偿的经典分歧。"
                    )
                elif dtype == "disagree_mild":
                    engine_text = (
                        f"Stockfish和Lc0对此局面略有不同看法"
                        f"（SF={sf_score:+.1f} vs Lc0={lc0_score:+.1f}），"
                        f"局面存在动态因素。"
                    )
                else:
                    engine_text = (
                        f"Lc0神经网络比Stockfish更看好此局面"
                        f"（{lc0_score:+.1f} vs {sf_score:+.1f}）。"
                    )
                paragraphs.append(("engine", engine_text))

        # 残局库（仅在最后阶段输出）
        if seg_end >= total_moves - 6 and tb_results:
            for tb in tb_results:
                if seg_start_mv <= tb["move_number"] <= seg_end:
                    paragraphs.append(("tablebase",
                        f"残局库确认：第{tb['move_number']}步（{tb['piece_count']}子局面）— {tb['verdict']}。"))

        i = seg_end

    # ===== 总结 =====
    err_count = sum(1 for s in steps if s.get("quality") in ("错误", "大错"))
    good_cnt = sum(1 for s in steps if s.get("quality") == "非常好")
    turning = [s for s in steps if abs(s.get("score_diff", 0)) > 2.0]

    summary = "纵观全局，"
    if err_count == 0:
        summary += f"双方发挥稳定，有{good_cnt}步被评为好棋，整盘棋展现了高水平的对弈。"
    else:
        summary += f"失误较多的一方（共{err_count}次错误或大错）未能保持稳定发挥。"
        if turning:
            w = max(turning, key=lambda s: abs(s.get("score_diff", 0)))
            summary += (
                f" 本局的转折点出现在第{round_num(w['move_number'])}回合"
                f"{w['side']}的{w['move_san']}（评分变化{w['score_diff']:+.1f}）。"
            )
    if tb_results:
        final_tb = tb_results[-1]
        summary += f" 残局库确认：{final_tb['verdict']}。"

    paragraphs.append(("summary", summary))

    # ===== 完整 PGN =====
    full_pgn = " ".join(
        f"{round_num(i+1)}. {sans[i]}" if (i+1) % 2 == 1 else sans[i]
        for i in range(len(sans))
    )
    paragraphs.append(("pgn_header", "完整PGN："))
    paragraphs.append(("pgn", full_pgn + f" {result_str}"))

    return title, paragraphs


# ===================== Word 文档生成 =====================

def generate_report(output_path=None, no_images=False):
    script_dir = Path(__file__).parent

    print("加载数据...")
    sans, headers, steps, phases, tb_results, step_comms = load_all(script_dir)

    title, paragraphs = build_article(sans, headers, steps, phases, tb_results, step_comms)
    print(f"文章构建完成: {len(paragraphs)} 段")

    # ---- 创建文档 ----
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(18)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # 标题区
    add_p(doc, "深蓝棋评", size=14, bold=True, color=C_TITLE,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=6, indent=False)
    add_p(doc, title, size=18, bold=True, color=C_TITLE,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=18, indent=False)

    # 初始化棋盘渲染器
    pieces_dir = script_dir / "pieces"
    board_renderer = BoardImageRenderer(pieces_dir) if not no_images else None
    if board_renderer and board_renderer.piece_imgs:
        print(f"棋盘渲染器已就绪（{len(board_renderer.piece_imgs)} 枚棋子）")
    tmpdir = None
    if board_renderer:
        tmpdir = Path(tempfile.mkdtemp(prefix="chess_boards_"))

    # 正文
    board_img_idx = 0
    for item in paragraphs:
        ptype = item[0]
        ptext = item[1]
        seg_end_mv = item[2] if len(item) > 2 else None
        if ptype in ("intro",):
            add_p(doc, ptext, size=11, after=12)
        elif ptype in ("body",):
            add_p(doc, ptext, size=11, after=6)
            # 为该段生成棋盘图
            if board_renderer and seg_end_mv and sans:
                board, last_fr, last_to = replay_to_move(sans, seg_end_mv)
                try:
                    # 检查是否将军
                    hl_sqs = []
                    if board.is_check():
                        king_sq = board.king(board.turn)
                        if king_sq is not None:
                            hl_sqs.append(king_sq)
                    # 着法标注
                    ann = None
                    if seg_end_mv <= len(steps):
                        q = steps[seg_end_mv - 1].get("quality", "")
                        ann = {"有疑问": "?!", "错误": "?", "大错": "??"}.get(q)
                    board_img = board_renderer.render(
                        board,
                        last_from=last_fr,
                        last_to=last_to,
                        highlight_squares=hl_sqs if hl_sqs else None,
                        annotation=ann,
                    )
                    img_path = tmpdir / f"board_{board_img_idx:03d}.png"
                    board_img = board_img.convert('RGB')
                    board_img.save(str(img_path), 'PNG')
                    board_img_idx += 1

                    # 插入图片
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_para.paragraph_format.space_after = Pt(10)
                    img_para.paragraph_format.space_before = Pt(2)
                    run = img_para.add_run()
                    run.add_picture(str(img_path), width=Cm(7.5))
                    # 图片标题
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_after = Pt(8)
                    cap_run = cap.add_run(
                        f"第{round_num(seg_end_mv)}回合走完后的局面")
                    cap_run.font.size = Pt(8)
                    cap_run.font.name = FONT
                    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
                    cap_run.font.color.rgb = C_FOOTER
                except Exception as e:
                    print(f"  ⚠ 棋盘图生成失败(第{seg_end_mv}步): {e}")
        elif ptype in ("critical",):
            add_p(doc, "⚠️ " + ptext, size=10.5, bold=True, color=C_RED, after=6)
        elif ptype in ("brilliant",):
            add_p(doc, "⭐ " + ptext, size=10.5, bold=True, after=6)
        elif ptype in ("tactics",):
            add_p(doc, "🎯 " + ptext, size=10.5, after=6)
        elif ptype in ("engine",):
            add_p(doc, "🔬 " + ptext, size=10, color=C_ENGINE, after=6)
        elif ptype in ("tablebase",):
            add_p(doc, "📚 " + ptext, size=10, bold=True, after=6)
        elif ptype in ("summary",):
            sep(doc)
            add_p(doc, ptext, size=11, bold=True, after=12)
        elif ptype in ("pgn_header",):
            sep(doc)
            add_p(doc, ptext, size=10, bold=True, color=C_FOOTER, after=4, indent=False)
        elif ptype in ("pgn",):
            add_p(doc, ptext, size=9, color=C_FOOTER, after=12, indent=False)

    # 页脚
    sep(doc)
    add_p(doc,
        f"本文由深蓝国际象棋协会AI系统自动生成，基于Stockfish + Lc0双引擎分析和Syzygy残局库查询。\n"
        f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}  |  仅供学习交流",
        size=8, color=C_FOOTER, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, indent=False)

    # 保存
    if output_path is None:
        sw = re.sub(r'[^\w]', '', headers.get("White", "w"))[:15]
        sb = re.sub(r'[^\w]', '', headers.get("Black", "b"))[:15]
        se = (headers.get("ECO", "?")).replace("/", "-")
        output_path = script_dir / f"chess_analysis_report_{sw}_vs_{sb}_{se}.docx"

    doc.save(str(output_path))
    print(f"\n棋评文档已生成: {output_path}")
    print(f"  段落: {len(paragraphs) + 2}")
    if board_renderer and tmpdir:
        print(f"  棋盘图: {board_img_idx} 张")
    # 清理临时文件
    if tmpdir and tmpdir.exists():
        import shutil
        shutil.rmtree(tmpdir, ignore_errors=True)
    return output_path


def main():
    import argparse
    ap = argparse.ArgumentParser(description="生成微信公众号风格的棋评 Word 文档")
    ap.add_argument("--output", type=str, default=None)
    ap.add_argument("--no-images", action="store_true", help="不嵌入棋盘图片")
    args = ap.parse_args()

    print("=" * 50)
    print("深蓝棋评 Word 文档生成器 v4（微信公众号风格 + 棋盘图）")
    print("=" * 50)
    generate_report(
        output_path=Path(args.output) if args.output else None,
        no_images=args.no_images,
    )


if __name__ == "__main__":
    main()
